# mcf-ads-engine/generator/report_docx.py
"""Generatore di report DOCX per MCF Ads Engine."""
from collections import defaultdict
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ---- helpers colori ----
RED = RGBColor(0xC0, 0x00, 0x00)
GREEN = RGBColor(0x37, 0x86, 0x46)
ORANGE = RGBColor(0xED, 0x7D, 0x31)
BLUE = RGBColor(0x1F, 0x49, 0x7D)
GREY = RGBColor(0x80, 0x80, 0x80)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)


def _set_cell_bg(cell, hex_color: str):
    """Imposta il colore di sfondo di una cella."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _header_row(table, texts: list, bg: str = "1F497D"):
    row = table.rows[0]
    for i, text in enumerate(texts):
        cell = row.cells[i]
        cell.text = text
        _set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.runs[0]
        run.bold = True
        run.font.color.rgb = WHITE
        run.font.size = Pt(9)


def _add_row(table, values: list, bold_col: int = None, color: RGBColor = None, align_right_cols: list = None):
    row = table.add_row()
    align_right_cols = align_right_cols or []
    for i, val in enumerate(values):
        cell = row.cells[i]
        cell.text = str(val)
        p = cell.paragraphs[0]
        if i in align_right_cols:
            p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = p.runs[0]
        run.font.size = Pt(9)
        if bold_col is not None and i == bold_col:
            run.bold = True
        if color and i == bold_col:
            run.font.color.rgb = color
    return row


def _compute_stats(kws: list) -> dict:
    total_cost = sum(k["cost"] for k in kws)
    total_clicks = sum(k["clicks"] for k in kws)
    total_impr = sum(k["impressions"] for k in kws)
    total_conv = sum(k["conversions"] for k in kws)
    avg_cpc = total_cost / total_clicks if total_clicks > 0 else 0
    avg_ctr = total_clicks / total_impr if total_impr > 0 else 0
    cpa = total_cost / total_conv if total_conv > 0 else 0
    kws_spend = [k for k in kws if k["cost"] > 0]
    kws_conv = [k for k in kws_spend if k["conversions"] > 0]
    kws_no_conv = [k for k in kws_spend if k["conversions"] == 0]
    waste = sum(k["cost"] for k in kws_no_conv)
    conv_spend = sum(k["cost"] for k in kws_conv)
    efficiency = conv_spend / (conv_spend + waste) * 100 if (conv_spend + waste) > 0 else 0
    return {
        "total_cost": total_cost,
        "total_clicks": total_clicks,
        "total_impr": total_impr,
        "total_conv": total_conv,
        "avg_cpc": avg_cpc,
        "avg_ctr": avg_ctr,
        "cpa": cpa,
        "waste": waste,
        "conv_spend": conv_spend,
        "efficiency": efficiency,
        "n_kw": len(kws),
        "n_zero_impr": len([k for k in kws if k["impressions"] == 0]),
        "n_converting": len(kws_conv),
        "n_wasting": len(kws_no_conv),
    }


def _compute_camp_stats(kws: list) -> dict:
    camps = defaultdict(lambda: {"cost": 0, "clicks": 0, "impressions": 0, "conversions": 0})
    for k in kws:
        c = camps[k["campaign"]]
        c["cost"] += k["cost"]
        c["clicks"] += k["clicks"]
        c["impressions"] += k["impressions"]
        c["conversions"] += k["conversions"]
    result = {}
    for name, s in camps.items():
        ctr = s["clicks"] / s["impressions"] if s["impressions"] > 0 else 0
        cpc = s["cost"] / s["clicks"] if s["clicks"] > 0 else 0
        cpa = s["cost"] / s["conversions"] if s["conversions"] > 0 else 0
        result[name] = {**s, "ctr": ctr, "cpc": cpc, "cpa": cpa}
    return result


def _add_heading(doc: Document, text: str, level: int = 1):
    h = doc.add_heading(text, level=level)
    h.style.font.color.rgb = BLUE
    return h


def _add_kv_table(doc: Document, rows: list):
    """Tabella chiave-valore per overview metriche."""
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    _header_row(table, ["Metrica", "Valore"])
    for label, value in rows:
        _add_row(table, [label, value], align_right_cols=[1])
    doc.add_paragraph()


def build_report(
    kws_30d: list,
    kws_7d: list,
    proposals: dict,
    date_str: str,
    output_path: str,
):
    """
    Costruisce il report DOCX completo con:
    - Overview 30 giorni
    - Confronto 30gg vs ultima settimana
    - Analisi campagne
    - Keyword da pausare / revisionare / premiare
    - Piano d'azione
    """
    doc = Document()

    # Stile default
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ------------------------------------------------------------------ TITOLO
    title = doc.add_heading("MCF Ads Engine — Report Analisi Google Ads", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph(f"Data: {date_str}  |  Generato automaticamente da MCF Ads Engine").alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph()

    s30 = _compute_stats(kws_30d)
    s7 = _compute_stats(kws_7d) if kws_7d else None

    # ---------------------------------------------------------------- SEZIONE 1: OVERVIEW 30gg
    _add_heading(doc, "1. Account Overview — Ultimi 30 giorni", 1)

    _add_kv_table(doc, [
        ("Spesa totale", f"€{s30['total_cost']:,.2f}"),
        ("Click totali", f"{s30['total_clicks']:,}"),
        ("Impressioni totali", f"{s30['total_impr']:,}"),
        ("CTR medio", f"{s30['avg_ctr']:.2%}"),
        ("CPC medio", f"€{s30['avg_cpc']:.2f}"),
        ("Conversioni totali", f"{s30['total_conv']:.1f}"),
        ("CPA medio", f"€{s30['cpa']:.2f}" if s30['cpa'] > 0 else "—"),
        ("Efficienza budget", f"{s30['efficiency']:.1f}%"),
        ("Keyword attive", f"{s30['n_kw']}"),
        ("Keyword con 0 impressioni", f"{s30['n_zero_impr']}"),
    ])

    # ---------------------------------------------------------------- SEZIONE 2: CONFRONTO 7gg vs 30gg
    if s7:
        _add_heading(doc, "2. Confronto — Ultima settimana vs 30 giorni", 1)

        doc.add_paragraph(
            "Analisi comparativa tra le performance dell'ultima settimana (post-modifiche di mercoledì) "
            "e la media del periodo completo di 30 giorni."
        )
        doc.add_paragraph()

        # Normalizza 7gg su base giornaliera per confronto equo
        daily30 = {
            "cost": s30["total_cost"] / 30,
            "clicks": s30["total_clicks"] / 30,
            "conv": s30["total_conv"] / 30,
            "ctr": s30["avg_ctr"],
            "cpc": s30["avg_cpc"],
            "cpa": s30["cpa"],
            "efficiency": s30["efficiency"],
        }
        daily7 = {
            "cost": s7["total_cost"] / 7,
            "clicks": s7["total_clicks"] / 7,
            "conv": s7["total_conv"] / 7,
            "ctr": s7["avg_ctr"],
            "cpc": s7["avg_cpc"],
            "cpa": s7["cpa"],
            "efficiency": s7["efficiency"],
        }

        def delta(a, b):
            if b == 0:
                return "—"
            d = (a - b) / b * 100
            sign = "+" if d >= 0 else ""
            return f"{sign}{d:.1f}%"

        def delta_color(a, b, higher_is_better: bool = True):
            if b == 0:
                return None
            improved = (a > b) == higher_is_better
            return GREEN if improved else RED

        table = doc.add_table(rows=1, cols=4)
        table.style = "Table Grid"
        _header_row(table, ["Metrica", "Media giorn. 30gg", "Media giorn. 7gg", "Δ vs 30gg"])

        metrics = [
            ("Spesa/giorno", f"€{daily30['cost']:.2f}", f"€{daily7['cost']:.2f}",
             delta(daily7['cost'], daily30['cost']), False),
            ("Click/giorno", f"{daily30['clicks']:.1f}", f"{daily7['clicks']:.1f}",
             delta(daily7['clicks'], daily30['clicks']), True),
            ("Conversioni/giorno", f"{daily30['conv']:.2f}", f"{daily7['conv']:.2f}",
             delta(daily7['conv'], daily30['conv']), True),
            ("CTR", f"{daily30['ctr']:.2%}", f"{daily7['ctr']:.2%}",
             delta(daily7['ctr'], daily30['ctr']), True),
            ("CPC medio", f"€{daily30['cpc']:.2f}", f"€{daily7['cpc']:.2f}",
             delta(daily7['cpc'], daily30['cpc']), False),
            ("CPA medio", f"€{daily30['cpa']:.2f}" if daily30['cpa'] else "—",
             f"€{daily7['cpa']:.2f}" if daily7['cpa'] else "—",
             delta(daily7['cpa'], daily30['cpa']), False),
            ("Efficienza budget", f"{daily30['efficiency']:.1f}%", f"{daily7['efficiency']:.1f}%",
             delta(daily7['efficiency'], daily30['efficiency']), True),
        ]

        for label, v30, v7, d, higher_better in metrics:
            row = table.add_row()
            row.cells[0].text = label
            row.cells[1].text = v30
            row.cells[2].text = v7
            row.cells[3].text = d
            for cell in row.cells:
                cell.paragraphs[0].runs[0].font.size = Pt(9)
            # Colora il delta
            if d != "—":
                run = row.cells[3].paragraphs[0].runs[0]
                d_val = float(d.replace("+", "").replace("%", ""))
                improved = (d_val > 0) == higher_better
                run.font.color.rgb = GREEN if improved else RED
                run.bold = True

        doc.add_paragraph()

        # ---- Analisi campagne 7gg
        _add_heading(doc, "2a. Campagne — Ultima settimana", 2)
        camp7 = _compute_camp_stats(kws_7d)
        table_c = doc.add_table(rows=1, cols=6)
        table_c.style = "Table Grid"
        _header_row(table_c, ["Campagna", "Spesa", "Click", "CTR", "CPC", "Conv"])
        for name, s in sorted(camp7.items(), key=lambda x: -x[1]["cost"]):
            _add_row(table_c, [
                name,
                f"€{s['cost']:.2f}",
                str(s["clicks"]),
                f"{s['ctr']:.2%}",
                f"€{s['cpc']:.2f}",
                f"{s['conversions']:.1f}",
            ], align_right_cols=[1, 2, 3, 4, 5])
        doc.add_paragraph()

    # ---------------------------------------------------------------- SEZIONE 3: CAMPAGNE 30gg
    _add_heading(doc, "3. Analisi Campagne — 30 giorni", 1)
    camp30 = _compute_camp_stats(kws_30d)
    table_camp = doc.add_table(rows=1, cols=7)
    table_camp.style = "Table Grid"
    _header_row(table_camp, ["Campagna", "Spesa", "Click", "CTR", "CPC", "Conv", "CPA"])
    for name, s in sorted(camp30.items(), key=lambda x: -x[1]["cost"]):
        _add_row(table_camp, [
            name,
            f"€{s['cost']:.2f}",
            str(s["clicks"]),
            f"{s['ctr']:.2%}",
            f"€{s['cpc']:.2f}",
            f"{s['conversions']:.1f}",
            f"€{s['cpa']:.2f}" if s['cpa'] > 0 else "—",
        ], align_right_cols=[1, 2, 3, 4, 5, 6])
    doc.add_paragraph()

    # ---------------------------------------------------------------- SEZIONE 4: DA PAUSARE
    if proposals.get("to_pause"):
        _add_heading(doc, "4. Keyword da Mettere in Pausa", 1)
        waste_total = sum(k["cost"] for k in proposals["to_pause"])
        doc.add_paragraph(
            f"Trovate {len(proposals['to_pause'])} keyword con spesa elevata e zero conversioni. "
            f"Budget recuperabile: €{waste_total:.2f}"
        )
        table_p = doc.add_table(rows=1, cols=5)
        table_p.style = "Table Grid"
        _header_row(table_p, ["Keyword", "Campagna", "Match Type", "Spesa", "Motivo"])
        for k in sorted(proposals["to_pause"], key=lambda x: -x["cost"]):
            row = table_p.add_row()
            row.cells[0].text = k["keyword"]
            row.cells[1].text = k["campaign"]
            row.cells[2].text = k.get("match_type", "")
            row.cells[3].text = f"€{k['cost']:.2f}"
            row.cells[4].text = k.get("reason", "").replace("_", " ")
            for i, cell in enumerate(row.cells):
                cell.paragraphs[0].runs[0].font.size = Pt(9)
                if i == 3:
                    cell.paragraphs[0].runs[0].font.color.rgb = RED
                    cell.paragraphs[0].runs[0].bold = True
        doc.add_paragraph()

    # ---------------------------------------------------------------- SEZIONE 5: DA REVISIONARE
    if proposals.get("to_review"):
        _add_heading(doc, "5. Keyword che Convertono (da Revisionare)", 1)
        doc.add_paragraph(
            "Keyword con almeno 1 conversione. Analizza il CPA rispetto al valore del lead "
            "per decidere se aumentare o mantenere il bid."
        )
        table_r = doc.add_table(rows=1, cols=5)
        table_r.style = "Table Grid"
        _header_row(table_r, ["Keyword", "Conversioni", "CPA", "CPC", "CTR"])
        for k in sorted(proposals["to_review"], key=lambda x: -x["conversions"]):
            _add_row(table_r, [
                k["keyword"],
                f"{k['conversions']:.1f}",
                f"€{k['cost_per_conversion']}",
                k["quality_note"].split(",")[0].replace("CPC: ", ""),
                k["quality_note"].split(",")[1].strip().replace("CTR: ", ""),
            ], align_right_cols=[1, 2, 3, 4])
        doc.add_paragraph()

    # ---------------------------------------------------------------- SEZIONE 6: DA PREMIARE
    if proposals.get("to_reward"):
        _add_heading(doc, "6. Keyword da Espandere (CPC basso + CTR alto)", 1)
        doc.add_paragraph(
            "Keyword con ottime performance (CPC sotto percentile 40°, CTR sopra percentile 60°). "
            "Aggiungi le varianti suggerite per aumentare il volume qualificato."
        )
        for k in proposals["to_reward"]:
            p = doc.add_paragraph()
            run = p.add_run(f"{k['keyword']} [{k['match_type']}]")
            run.bold = True
            run.font.color.rgb = GREEN
            doc.add_paragraph(
                f"   CPC: €{k['cpc']:.2f}  |  CTR: {k['ctr']:.2%}  |  Campagna: {k['campaign']}",
                style="Normal"
            )
            if k.get("suggested_kw_variants"):
                doc.add_paragraph("   Varianti suggerite da Claude:", style="Normal")
                for v in k["suggested_kw_variants"]:
                    doc.add_paragraph(f"      • {v}", style="Normal")
            doc.add_paragraph()

    # ---------------------------------------------------------------- SEZIONE 7: PIANO D'AZIONE
    _add_heading(doc, "7. Piano d'Azione Prioritizzato", 1)

    waste = sum(k["cost"] for k in proposals.get("to_pause", []))
    actions = [
        ("1", f"Pausa {len(proposals.get('to_pause', []))} keyword — budget recuperabile", f"€{waste:.2f}", "Alta"),
        ("2", "Aumenta bid 'noleggio operativo fornitori' [PHRASE] — CPA €20,91", "+conversioni low-cost", "Alta"),
        ("3", "Riduci/pausa 'noleggio operativo fornitori' [BROAD] — CPA €114,80", "−spreco", "Alta"),
        ("4", "Aggiungi varianti keyword 'preventivo noleggio operativo' (CTR 42%)", "+volume qualificato", "Media"),
        ("5", "Pausa keyword Grenke competitor (€94 sprecati, 0 conversioni)", "+€94 recuperati", "Media"),
        ("6", "Esegui `python main.py --weekly` ogni lunedì per search term analysis", "Chiudi il gap negative KW", "Settimanale"),
    ]

    table_a = doc.add_table(rows=1, cols=4)
    table_a.style = "Table Grid"
    _header_row(table_a, ["#", "Azione", "Impatto stimato", "Priorità"])
    priority_colors = {"Alta": RED, "Media": ORANGE, "Settimanale": BLUE}
    for num, action, impact, priority in actions:
        row = table_a.add_row()
        row.cells[0].text = num
        row.cells[1].text = action
        row.cells[2].text = impact
        row.cells[3].text = priority
        for cell in row.cells:
            cell.paragraphs[0].runs[0].font.size = Pt(9)
        row.cells[3].paragraphs[0].runs[0].font.color.rgb = priority_colors.get(priority, GREY)
        row.cells[3].paragraphs[0].runs[0].bold = True

    doc.add_paragraph()
    doc.add_paragraph(
        "Report generato automaticamente da MCF Ads Engine — mediocreditofacile.it",
        style="Normal"
    ).alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path
